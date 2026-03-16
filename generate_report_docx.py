from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Glass Aero Production Tracking System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Development Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')

info = [
    ('Project:', 'Aviation Windshield Repair/Overhaul Production Tracking System'),
    ('Client:', 'Glass Aero (formerly Avemar Group)'),
    ('Development Period:', 'February 23, 2026 \u2013 March 12, 2026 (18 days)'),
    ('Total Commits:', '28'),
    ('Status:', 'Active Development'),
]
for label, value in info:
    p = doc.add_paragraph()
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    p.add_run(value)

# --- Section 1 ---
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The Glass Aero Production Tracking System is a web-based application built to manage the '
    'end-to-end workflow of aviation windshield repair and overhaul operations. The system tracks '
    'repair orders through an 18-stage production process, manages parts inventory and bills of '
    'materials, provides real-time dashboards, and generates operational reports. The application '
    'is hosted on GitHub Pages with a self-hosted Supabase (PostgreSQL) backend.'
)

# --- Section 2 ---
doc.add_heading('2. Technology Stack', level=1)
tech_data = [
    ('Component', 'Technology'),
    ('Frontend', 'HTML, vanilla JavaScript, Tailwind CSS (CDN)'),
    ('Backend/DB', 'Supabase (PostgreSQL) \u2014 self-hosted'),
    ('Authentication', 'Supabase Auth with role-based access control'),
    ('Hosting', 'GitHub Pages'),
    ('Version Control', 'Git / GitHub'),
]
table = doc.add_table(rows=len(tech_data), cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (col1, col2) in enumerate(tech_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for cell in table.rows[i].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

# --- Section 3 ---
doc.add_heading('3. Application Pages', level=1)
pages_data = [
    ('Page', 'Purpose'),
    ('Dashboard', 'Production pipeline overview with stage cards'),
    ('Login', 'Secure authentication'),
    ('Repair Orders', 'Searchable/filterable list of all repair orders'),
    ('New Repair Order', 'Create new windshield repair work orders'),
    ('Repair Order Detail', 'Single order view with stage progress and actions'),
    ('Inventory', 'Subcomponent parts inventory management'),
    ('Bill of Materials (BOM)', 'BOM configuration per windshield part number'),
    ('Reports', 'Operational and production reports'),
    ('Import', 'Bulk import of repair orders via Excel upload'),
    ('Settings', 'Admin configuration (parts, stages, hours)'),
]
table = doc.add_table(rows=len(pages_data), cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (col1, col2) in enumerate(pages_data):
    table.rows[i].cells[0].text = col1
    table.rows[i].cells[1].text = col2
    if i == 0:
        for cell in table.rows[i].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

# --- Section 4 ---
doc.add_heading('4. Database Schema', level=1)
doc.add_paragraph('The system uses 10 database tables:')
db_tables = [
    'production_parts \u2014 Windshield part numbers (King Air and CRJ variants)',
    'production_stages \u2014 18-step production workflow definition',
    'repair_orders \u2014 Core repair order records with status tracking',
    'stage_history \u2014 Audit trail of stage transitions with timestamps',
    'subcomponents \u2014 Repair materials inventory (parts, quantities, reorder points)',
    'bom_items \u2014 Bill of materials linking windshield parts to subcomponents per stage',
    'parts_issuance \u2014 Tracks parts issued to specific repair orders',
    'repair_order_documents \u2014 File uploads linked to repair orders',
    'hold_history \u2014 Records of work stoppages with reasons and durations',
    'app_settings \u2014 Configurable application settings (business hours, etc.)',
]
for t in db_tables:
    doc.add_paragraph(t, style='List Number')

# --- Section 5 ---
doc.add_heading('5. Development Timeline', level=1)

doc.add_heading('Phase 1: Foundation (February 23, 2026)', level=2)
phase1 = [
    'Initial project setup \u2014 Created repository, established project structure',
    'Core pages built \u2014 Dashboard, login, repair orders list, new repair order form, repair order detail, inventory, BOM, and settings pages',
    'Database schema \u2014 Designed and deployed 8 core tables with Row Level Security',
    'Authentication \u2014 Implemented Supabase Auth with session management',
    'Dashboard \u2014 Production pipeline visualization with stage cards showing unit counts',
    'Stage progression \u2014 Ability to advance repair orders through production stages',
    'BOM auto-issuance \u2014 Automatic parts issuance based on bill of materials when advancing stages',
    'Document uploads \u2014 File attachment capability at any production stage',
    'Stage reversion \u2014 Ability to move repair orders back to previous stages with inventory reversal',
    'UI refinements \u2014 Fixed dropdown styling, dashboard scrolling, background colors',
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Phase 2: BOM and Reporting (February 23\u201327, 2026)', level=2)
phase2 = [
    'BOM bug fix \u2014 Corrected issue where parts were duplicated when advancing through stages without BOM items',
    'Reports page \u2014 Built comprehensive reporting with production throughput, inventory projections, and delivery tracking',
    'Inventory improvements \u2014 Changed low stock calculations to show projected shortages based on in-work units',
    'Invoice tracking \u2014 Added invoice number field to repair orders',
    'Historical data import \u2014 SQL scripts to import existing Vertex repair order data',
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Phase 3: Production Controls (March 3\u20134, 2026)', level=2)
phase3 = [
    'Hold/Resume feature \u2014 Added ability to place repair orders on hold with reason tracking, visual indicators (red pulsing), hold history, and dashboard hold counts',
    'Edit repair orders \u2014 Post-creation editing of all repair order fields including part number changes with automatic BOM reversal and re-issuance',
    'Part number change logic \u2014 Comprehensive flow: reverse all issued parts, restore inventory, fetch new BOM, re-issue for completed stages',
    'Edit windshield part numbers \u2014 Settings page capability to modify existing part configurations',
    'Business hours tracking \u2014 Stage duration calculations based on configurable business hours (9 AM\u20136 PM Mon\u2013Fri EST) instead of wall-clock time, with settings page for adjustment',
    'Inspection checklist \u2014 Inspector selects which subsequent stages are needed based on inspection results; system automatically skips unselected stages and adjusts inventory projections',
    'Delete repair orders \u2014 Admin function with full cascade cleanup (reverse parts, delete history, remove documents)',
    'Sequential RO numbers \u2014 Changed from random to sequential format (RO-YYYYMMDD-0001)',
    'Bulk import \u2014 Excel template download, file upload, validation, and batch import of repair orders',
    'Cache management \u2014 Added version query parameters to all script tags to prevent browser caching issues on GitHub Pages',
]
for item in phase3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Phase 4: Access Control and Rebranding (March 11\u201312, 2026)', level=2)
phase4 = [
    'Company rebranding \u2014 Changed all references from "Avemar Group" to "Glass Aero" across all pages, scripts, and documentation',
    'Role-based access control \u2014 Implemented admin/standard user roles using Supabase Auth app_metadata. Standard users see Dashboard, Repair Orders, Inventory, and Reports. Admin users see all pages plus BOM, Settings, Import, and Delete functions.',
    'New windshield part numbers \u2014 Added King Air (101-384025) and CRJ (NP139321, 601R33033) part number families',
    'Production workflow update \u2014 Replaced 15-stage process with new 18-stage workflow per customer requirements. Combined old Receiving, Document Verification, and Inspection into single Receiving stage. Combined Heater Installation and Outer Glass Installation into Interlayer/Heater/Sensor Install. Added new stages for Conductive Coating, P1 Autoclave, Fiber Glass, Retainer, Edge Sealant, Weather Sealant, and Final Pics.',
]
for item in phase4:
    doc.add_paragraph(item, style='List Bullet')

# --- Section 6 ---
doc.add_heading('6. Current Production Stages (18-Step Workflow)', level=1)
stages_data = [
    ('Stage', 'Name', 'Role', 'Time Limit'),
    ('1', 'Receiving Inspection', 'Receiving', '2 hrs'),
    ('2', 'Disassembly', 'Shop Floor', '24 hrs'),
    ('3', 'Removal of Conductive Coating', 'Shop Floor', '24 hrs'),
    ('4', 'P1 Autoclave (skippable)', 'Shop Floor', '48 hrs'),
    ('5', 'Cleaning PRE-CAT3/4', 'Shop Floor', '16 hrs'),
    ('6', 'Interlayer, Heater, Sensor Install', 'Shop Floor', '24 hrs'),
    ('7', 'Autoclave', 'Shop Floor', '48 hrs'),
    ('8', 'Testing', 'Quality', '24 hrs'),
    ('9', 'Cleaning PRE-Fiber Glass', 'Shop Floor', '16 hrs'),
    ('10', 'Fiber Glass Installation', 'Shop Floor', '24 hrs'),
    ('11', 'Retainer Installation', 'Shop Floor', '24 hrs'),
    ('12', 'Polishing', 'Shop Floor', '24 hrs'),
    ('13', 'Peripheral Edge Sealant PRC', 'Shop Floor', '24 hrs'),
    ('14', 'Weather Sealant PRC', 'Shop Floor', '24 hrs'),
    ('15', 'Cleaning', 'Shop Floor', '16 hrs'),
    ('16', 'Final Pics', 'Quality', '8 hrs'),
    ('17', 'Final Inspection', 'Quality', '16 hrs'),
    ('18', 'Shipping', 'Receiving', '8 hrs'),
]
table = doc.add_table(rows=len(stages_data), cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(stages_data):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

# --- Section 7 ---
doc.add_heading('7. Supported Windshield Part Numbers', level=1)

doc.add_heading('King Air Series (101-384025)', level=2)
doc.add_paragraph('101-384025-21 through -24 \u2014 King Air Windshield Assemblies', style='List Bullet')

doc.add_heading('CRJ Main Windshield (NP139321)', level=2)
doc.add_paragraph('NP139321-1, -2, -5, -6, -9 through -18', style='List Bullet')
doc.add_paragraph('Odd dash numbers = Left Hand (LH), Even = Right Hand (RH)', style='List Bullet')

doc.add_heading('CRJ Side Windows (601R33033)', level=2)
doc.add_paragraph('601R33033-3, -4, -11, -12, -19, -20, -23, -24, -29, -30', style='List Bullet')
doc.add_paragraph('Odd dash numbers = Left Hand (LH), Even = Right Hand (RH)', style='List Bullet')

# --- Section 8 ---
doc.add_heading('8. Key Features Summary', level=1)
features = [
    'Real-time dashboard with production pipeline visualization',
    '18-stage production workflow with configurable time limits',
    'Inspection checklist allowing inspectors to customize the repair path per unit',
    'Automatic stage skipping based on inspection results',
    'Bill of Materials with automatic parts issuance and inventory deduction',
    'Projected inventory calculations accounting for all in-work units',
    'Hold/Resume with reason tracking and visual indicators',
    'Business hours time tracking (configurable work schedule and timezone)',
    'Document attachments at any production stage',
    'Bulk import via Excel template',
    'Role-based access control (Admin vs Standard users)',
    'Sequential RO numbering (RO-YYYYMMDD-XXXX)',
    'Edit and delete repair orders with full inventory reversal',
    'Reports with production throughput, inventory projections, and delivery metrics',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# --- Section 9 ---
doc.add_heading('9. File Structure', level=1)
file_structure = """Glass Aero Production Tracking/
\u251c\u2500\u2500 index.html                  (Dashboard)
\u251c\u2500\u2500 login.html                  (Authentication)
\u251c\u2500\u2500 repair-orders.html          (Repair order list)
\u251c\u2500\u2500 new-repair-order.html       (Create repair order)
\u251c\u2500\u2500 repair-order-detail.html    (Single order detail)
\u251c\u2500\u2500 inventory.html              (Parts inventory)
\u251c\u2500\u2500 bom.html                    (Bill of materials)
\u251c\u2500\u2500 reports.html                (Reports)
\u251c\u2500\u2500 import.html                 (Bulk import)
\u251c\u2500\u2500 settings.html               (Admin settings)
\u251c\u2500\u2500 js/
\u2502   \u251c\u2500\u2500 supabase-config.js      (Database layer)
\u2502   \u251c\u2500\u2500 auth.js                 (Authentication and RBAC)
\u2502   \u251c\u2500\u2500 repair-order-detail.js  (Order detail logic)
\u2502   \u251c\u2500\u2500 repair-orders.js        (Order list logic)
\u2502   \u251c\u2500\u2500 new-repair-order.js     (New order form logic)
\u2502   \u251c\u2500\u2500 inventory.js            (Inventory page logic)
\u2502   \u251c\u2500\u2500 bom.js                  (BOM page logic)
\u2502   \u251c\u2500\u2500 reports.js              (Reports logic)
\u2502   \u251c\u2500\u2500 import.js               (Import logic)
\u2502   \u2514\u2500\u2500 settings.js             (Settings logic)
\u2514\u2500\u2500 sql/
    \u251c\u2500\u2500 schema.sql              (Core database schema)
    \u251c\u2500\u2500 add-hold-feature.sql    (Hold/resume migration)
    \u251c\u2500\u2500 add-business-hours.sql  (Business hours settings)
    \u251c\u2500\u2500 add-inspection-checklist.sql
    \u251c\u2500\u2500 update-stages-v2.sql    (18-stage workflow)
    \u2514\u2500\u2500 add-admin-role.md       (Admin role setup guide)"""
p = doc.add_paragraph()
run = p.add_run(file_structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# --- Section 10 ---
doc.add_heading('10. Development Effort', level=1)

doc.add_heading('Working Sessions (from git commit timestamps)', level=2)
sessions = [
    ('Date', 'Time Range', 'Session Duration'),
    ('Feb 23', '4:56 PM \u2013 9:25 PM', '~4.5 hours'),
    ('Feb 27', '1:20 PM \u2013 7:04 PM', '~5.75 hours'),
    ('Mar 3', '8:19 PM \u2013 8:48 PM', '~0.5 hours'),
    ('Mar 4', '6:16 PM \u2013 8:20 PM', '~2 hours'),
    ('Mar 11', '7:44 AM \u2013 8:12 AM', '~0.5 hours'),
    ('Mar 12', '9:27 PM \u2013 9:39 PM', '~0.25 hours'),
]
table = doc.add_table(rows=len(sessions), cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(sessions):
    for j, val in enumerate(row_data):
        table.rows[i].cells[j].text = val
        if i == 0:
            for run in table.rows[i].cells[j].paragraphs[0].runs:
                run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Tracked commit time: ~13.5 hours')
run.bold = True

doc.add_paragraph(
    'Commits capture only when code was pushed. Additional time was spent on planning, '
    'requirements discussion, testing, troubleshooting, Supabase SQL migrations, and UI verification '
    'between commits.'
)

p = doc.add_paragraph()
run = p.add_run('Estimated total development effort: 20\u201325 hours')
run.bold = True

# --- Section 11 ---
doc.add_heading('11. Deployment', level=1)
deploy = [
    ('Repository:', 'GitHub (kc10guru/Avemar-Production-Tracking)'),
    ('Live URL:', 'https://kc10guru.github.io/Avemar-Production-Tracking'),
    ('Database:', 'Self-hosted Supabase at https://avemar-db.duckdns.org'),
    ('Updates:', 'Push to main branch triggers automatic GitHub Pages deployment'),
]
for label, value in deploy:
    p = doc.add_paragraph()
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    p.add_run(value)

doc.add_paragraph('')
p = doc.add_paragraph('Report generated March 12, 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.italic = True
p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output_path = r'c:\Users\kc10g\OneDrive\Documents\JCD Enterprises\Avemar-Production-Tracking\Glass Aero Production Tracker - Development Report.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
