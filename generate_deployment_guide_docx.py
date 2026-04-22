"""
Convert Glass Aero Production Tracker - On-Premises Deployment Guide.md to .docx
Requires: pip install python-docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


def strip_md_cell(s: str) -> str:
    s = s.strip()
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def add_paragraph_with_inline(para, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            para.add_run(part)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def flush_table(doc: Document, rows: list[str]) -> None:
    if not rows:
        return
    cells0 = [strip_md_cell(c) for c in rows[0].split("|")[1:-1]]
    ncols = len(cells0)
    if ncols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        parts = [strip_md_cell(c) for c in row.split("|")[1:-1]]
        for ci in range(ncols):
            cell_text = parts[ci] if ci < len(parts) else ""
            table.rows[ri].cells[ci].text = cell_text
            if ri == 0:
                for paragraph in table.rows[ri].cells[ci].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    code_mode = False
    code_buf: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if code_mode:
                add_code_block(doc, code_buf)
                code_buf = []
                code_mode = False
            else:
                code_mode = True
            i += 1
            continue

        if code_mode:
            code_buf.append(raw)
            i += 1
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            table_rows: list[str] = []
            while i < len(lines):
                s = lines[i].strip()
                if not s.startswith("|"):
                    break
                if re.match(r"^\|[\s\-:|]+\|$", s):
                    i += 1
                    continue
                table_rows.append(lines[i])
                i += 1
            flush_table(doc, table_rows)
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            ht = m.group(2).strip()
            ht = re.sub(r"\*\*(.+?)\*\*", r"\1", ht)
            doc.add_heading(ht, level=level)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_paragraph_with_inline(p, m.group(2).strip())
            i += 1
            continue

        m = re.match(r"^[-*]\s+(.+)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_paragraph_with_inline(p, m.group(1).strip())
            i += 1
            continue

        p = doc.add_paragraph()
        add_paragraph_with_inline(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent
    md_file = root / "Glass Aero Production Tracker - On-Premises Deployment Guide.md"
    out_file = root / "Glass Aero Production Tracker - On-Premises Deployment Guide.docx"
    if not md_file.is_file():
        raise SystemExit(f"Missing: {md_file}")
    convert(md_file, out_file)
