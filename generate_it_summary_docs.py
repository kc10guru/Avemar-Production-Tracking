from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_executive_summary():
    """Create a 1-2 page high-level overview for IT managers/stakeholders"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Glass Aero Production Tracker', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('IT Overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph('')
    
    # Section 1: Purpose
    doc.add_heading('What is the Production Tracker?', level=1)
    doc.add_paragraph(
        'A web-based application for tracking aviation windshield repair orders through '
        'an 18-stage production workflow. It manages inventory, tracks parts usage, '
        'provides real-time dashboards, and generates production reports.'
    )
    
    # Section 2: Data Flow (simple)
    doc.add_heading('Data Flow', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Data In: ').bold = True
    p.add_run('Repair orders (manual entry or Excel import), inventory receipts, document uploads')
    
    p = doc.add_paragraph()
    p.add_run('Data Out: ').bold = True
    p.add_run('Production dashboard, weekly/quarterly/annual reports, barcode labels, audit history')
    
    p = doc.add_paragraph()
    p.add_run('External Integrations: ').bold = True
    p.add_run('None – all data is entered manually or via Excel import')
    
    # Section 3: Architecture (simplified)
    doc.add_heading('System Architecture', level=1)
    
    arch_data = [
        ('Component', 'Technology', 'Notes'),
        ('Frontend', 'HTML / JavaScript', 'Static web pages, no server-side code'),
        ('Backend', 'Supabase (PostgreSQL)', 'Database, authentication, file storage'),
        ('Hosting', 'GitHub Pages', 'Can be moved to any web server'),
    ]
    table = doc.add_table(rows=len(arch_data), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(arch_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph(
        'The system consists of static web files that connect to a PostgreSQL database via API. '
        'No application server is required – the browser communicates directly with the database backend.'
    )
    
    # Section 4: Dependencies
    doc.add_heading('Dependencies & Offline Capability', level=1)
    
    deps = [
        'Internet connection required (database is network-accessible)',
        'Supabase backend (PostgreSQL + API + Auth) – currently self-hosted',
        'Web browser (Chrome, Edge, Firefox, Safari)',
    ]
    for d in deps:
        doc.add_paragraph(d, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Offline Mode: ').bold = True
    p.add_run('Not currently supported. Could be modified for full on-premises deployment if needed.')
    
    # Section 5: Authentication
    doc.add_heading('Authentication & Access Control', level=1)
    
    auth_data = [
        ('Method', 'Email + Password'),
        ('Provider', 'Supabase Auth (built-in)'),
        ('Roles', 'Admin and Standard User'),
    ]
    for label, value in auth_data:
        p = doc.add_paragraph()
        p.add_run(label + ': ').bold = True
        p.add_run(value)
    
    doc.add_paragraph('')
    
    access_data = [
        ('Capability', 'Standard', 'Admin'),
        ('View dashboard & orders', '✓', '✓'),
        ('Advance repair orders', '✓', '✓'),
        ('Manage inventory', '✓', '✓'),
        ('View reports', '✓', '✓'),
        ('Configure BOM', '–', '✓'),
        ('System settings', '–', '✓'),
        ('Import / Delete', '–', '✓'),
    ]
    table = doc.add_table(rows=len(access_data), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(access_data):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # Section 6: Deployment Options
    doc.add_heading('Deployment Options', level=1)
    
    doc.add_paragraph(
        'The system can be deployed in multiple configurations:'
    )
    
    deploy_options = [
        'Current: Frontend on GitHub Pages, Supabase self-hosted on developer server',
        'On-Premises: All components hosted on client infrastructure (ESXi/Docker)',
        'Cloud: Supabase cloud or any PostgreSQL provider + static file hosting',
    ]
    for opt in deploy_options:
        doc.add_paragraph(opt, style='List Bullet')
    
    # Footer
    doc.add_paragraph('')
    p = doc.add_paragraph('For detailed technical specifications, see the Technical Specification document.')
    p.runs[0].font.italic = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph('April 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    return doc


def create_technical_specification():
    """Create the detailed technical spec for IT staff who will deploy/maintain"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Glass Aero Production Tracker', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Specification')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph('')
    
    # --- Section 1: Purpose ---
    doc.add_heading('1. Application Purpose', level=1)
    doc.add_paragraph(
        'The Glass Aero Production Tracker is a web-based repair order management system '
        'for aviation windshield repair and overhaul operations. It tracks repair orders through '
        'an 18-stage production workflow, manages parts inventory with automatic Bill of Materials (BOM) '
        'issuance, provides real-time production dashboards, and generates operational reports.'
    )
    
    doc.add_heading('Business Problems Solved', level=2)
    problems = [
        'Track each windshield unit from receiving through shipping',
        'Ensure parts are issued at the correct production stage',
        'Identify late or held units in real-time',
        'Generate production throughput and inventory reports',
        'Maintain audit trail of all stage transitions',
    ]
    for p in problems:
        doc.add_paragraph(p, style='List Bullet')
    
    # --- Section 2: Data Ingestion ---
    doc.add_heading('2. Data Ingestion', level=1)
    
    data_in = [
        ('Data Type', 'Source', 'Method'),
        ('Repair Orders', 'Manual entry or Excel bulk import', 'Web form / .xlsx upload'),
        ('Inventory Levels', 'Manual entry via Receive Stock', 'Web form'),
        ('Stage Transitions', 'User action (Advance Stage)', 'Web interface'),
        ('Documents', 'User upload (PDFs, images)', 'Drag-and-drop / file picker'),
        ('Historical Data', 'One-time SQL import', 'SQL script'),
    ]
    table = doc.add_table(rows=len(data_in), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(data_in):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Note: ').bold = True
    p.add_run('No external system integrations. All data is entered manually or via Excel import.')
    
    # --- Section 3: System Outputs ---
    doc.add_heading('3. System Outputs', level=1)
    
    outputs = [
        ('Output', 'Description'),
        ('Dashboard', 'Real-time production pipeline showing units at each stage, late/hold counts, low stock alerts'),
        ('Weekly Report', 'Units received, delivered, in-work, behind schedule'),
        ('Quarterly Report', 'Production trends, monthly breakdown, average days to complete'),
        ('Annual Report', 'Year-over-year comparison, 12-month trends'),
        ('Projected Inventory', 'Current stock vs. projected needs based on in-work orders'),
        ('Barcode Labels', 'Printable labels with Code 128 barcode (Brother QL-1100C compatible)'),
        ('Stage History', 'Full audit trail of who advanced each stage and when'),
        ('Parts Issuance Log', 'Record of all inventory parts issued to each repair order'),
    ]
    table = doc.add_table(rows=len(outputs), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(outputs):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # --- Section 4: Dependencies ---
    doc.add_heading('4. External Dependencies', level=1)
    
    deps = [
        ('Dependency', 'Required?', 'Notes'),
        ('Internet', 'Yes (currently)', 'Frontend loads from GitHub Pages; JS library from CDN'),
        ('Supabase Backend', 'Yes', 'PostgreSQL database + Auth + Realtime + Storage'),
        ('GitHub Pages', 'Yes (currently)', 'Hosts static frontend files'),
        ('Tailwind CSS CDN', 'Yes (currently)', 'CSS framework loaded at runtime'),
        ('Supabase JS CDN', 'Yes (currently)', 'Database client library'),
    ]
    table = doc.add_table(rows=len(deps), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(deps):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    doc.add_heading('Offline Capability', level=2)
    doc.add_paragraph(
        'Not supported in current configuration. The frontend depends on CDN-hosted libraries '
        'and a network-accessible Supabase instance.'
    )
    doc.add_paragraph(
        'With modifications: The system could be made fully self-contained by bundling dependencies '
        'and hosting everything on-premises. The codebase is vanilla HTML/JS with no build step, '
        'making this straightforward.'
    )
    
    # --- Section 5: Current Setup ---
    doc.add_heading('5. Current Deployment', level=1)
    
    setup = [
        ('Component', 'Technology', 'Location'),
        ('Frontend', 'Static HTML + JavaScript + CSS', 'GitHub Pages'),
        ('Database', 'PostgreSQL (via Supabase)', 'Self-hosted at avemar-db.duckdns.org'),
        ('Authentication', 'Supabase Auth (GoTrue)', 'Same Supabase instance'),
        ('File Storage', 'Supabase Storage', 'Same Supabase instance'),
        ('Supabase Runtime', 'Docker containers', 'Developer server'),
    ]
    table = doc.add_table(rows=len(setup), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(setup):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    doc.add_heading('Supabase Docker Components', level=2)
    supabase_components = [
        'PostgreSQL – Primary database',
        'PostgREST – Auto-generated REST API',
        'GoTrue – Authentication service',
        'Realtime – WebSocket subscriptions for live updates',
        'Storage API – File uploads and downloads',
    ]
    for c in supabase_components:
        doc.add_paragraph(c, style='List Bullet')
    
    # --- Section 6: Architecture ---
    doc.add_heading('6. Architecture Diagram', level=1)
    
    diagram = """┌─────────────────────────────────────────────────────────────┐
│                        USER BROWSER                         │
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────────────┐  │
│  │ HTML Pages  │  │  JavaScript │  │ Tailwind CSS (CDN)  │  │
│  └──────┬──────┘  └──────┬──────┘  └─────────────────────┘  │
└─────────┼────────────────┼──────────────────────────────────┘
          │                │
          │     HTTPS      │ Supabase JS Client
          │                │
          ▼                ▼
┌─────────────────────────────────────────────────────────────┐
│               SELF-HOSTED SUPABASE (Docker)                 │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────────┐   │
│  │  PostgreSQL  │  │   PostgREST  │  │     GoTrue       │   │
│  │  (Database)  │  │  (REST API)  │  │ (Authentication) │   │
│  └──────────────┘  └──────────────┘  └──────────────────┘   │
│  ┌──────────────┐  ┌──────────────┐                         │
│  │   Realtime   │  │   Storage    │                         │
│  │ (WebSockets) │  │   (Files)    │                         │
│  └──────────────┘  └──────────────┘                         │
└─────────────────────────────────────────────────────────────┘"""
    
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    
    # --- Section 7: Database ---
    doc.add_heading('7. Database Schema', level=1)
    
    tables = [
        ('Table', 'Purpose'),
        ('production_parts', 'Windshield part number catalog'),
        ('production_stages', '18-step workflow definition'),
        ('repair_orders', 'Core repair order records'),
        ('stage_history', 'Audit trail of stage transitions'),
        ('subcomponents', 'Parts inventory (qty, reorder point, etc.)'),
        ('bom_items', 'Bill of materials per part number/stage'),
        ('parts_issuance', 'Parts issued to repair orders'),
        ('repair_order_documents', 'File upload metadata'),
        ('hold_history', 'Work stoppage records'),
        ('app_settings', 'Business hours and configuration'),
    ]
    table = doc.add_table(rows=len(tables), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(tables):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Security: ').bold = True
    p.add_run('Row Level Security (RLS) enabled on all tables. All API calls authenticated via JWT.')
    
    # --- Section 8: Authentication ---
    doc.add_heading('8. Authentication & Access Control', level=1)
    
    auth_info = [
        ('Auth Provider', 'Supabase Auth (GoTrue)'),
        ('Login Method', 'Email + Password'),
        ('Session', 'JWT token stored in browser localStorage'),
        ('Roles', 'Two roles: admin and standard'),
        ('Role Storage', 'Supabase Auth app_metadata.role field'),
    ]
    for label, value in auth_info:
        p = doc.add_paragraph()
        p.add_run(label + ': ').bold = True
        p.add_run(value)
    
    doc.add_paragraph('')
    doc.add_heading('Access Matrix', level=2)
    
    access = [
        ('Feature', 'Standard User', 'Admin'),
        ('Dashboard', '✓', '✓'),
        ('View/Advance Repair Orders', '✓', '✓'),
        ('Inventory (view/receive)', '✓', '✓'),
        ('Reports', '✓', '✓'),
        ('BOM Configuration', '✗', '✓'),
        ('Settings', '✗', '✓'),
        ('Import (bulk upload)', '✗', '✓'),
        ('Delete Repair Orders', '✗', '✓'),
    ]
    table = doc.add_table(rows=len(access), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(access):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # --- Section 9: Migration ---
    doc.add_heading('9. Migration to Client Infrastructure', level=1)
    
    doc.add_paragraph('To deploy on client infrastructure, the following steps are required:')
    
    migration_steps = [
        'Provision a server (physical or VM) capable of running Docker',
        'Install Docker and Docker Compose',
        'Deploy Supabase stack using official docker-compose template',
        'Import database schema (SQL files provided)',
        'Configure DNS or update hosts file for database URL',
        'Host frontend files on any web server (IIS, nginx, Apache) or keep on GitHub Pages',
        'Update supabase-config.js with new database URL and API key',
        'Create user accounts and assign admin roles',
        'Configure SSL certificates for HTTPS',
        'Set up database backups (pg_dump scheduled task)',
    ]
    for i, step in enumerate(migration_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # Footer
    doc.add_paragraph('')
    p = doc.add_paragraph('April 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    return doc


def create_migration_sow():
    """Create Migration Statement of Work document"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Statement of Work', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Glass Aero Production Tracker — On-Premises Migration')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph('')
    
    # Document info
    info = [
        ('Document Version:', '1.0'),
        ('Date:', 'April 2026'),
        ('Prepared By:', 'JCD Enterprises'),
        ('Prepared For:', 'Glass Aero'),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.add_run(label + ' ').bold = True
        p.add_run(value)
    
    # Section 1: Overview
    doc.add_heading('1. Project Overview', level=1)
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'This Statement of Work (SOW) defines the scope, deliverables, timeline, and responsibilities '
        'for migrating the Glass Aero Production Tracker from its current hosted environment to '
        'Glass Aero\'s internal on-premises infrastructure.'
    )
    
    doc.add_heading('1.2 Objectives', level=2)
    objectives = [
        'Deploy all application components on Glass Aero infrastructure',
        'Eliminate external dependencies (CDN, cloud services)',
        'Migrate existing data from current system',
        'Ensure system operates fully within internal network',
        'Transfer operational knowledge to Glass Aero IT staff',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # Section 2: Scope
    doc.add_heading('2. Scope of Work', level=1)
    
    doc.add_heading('2.1 In Scope', level=2)
    in_scope = [
        ('Infrastructure Setup', 'Deploy Docker environment with all required services'),
        ('Frontend Deployment', 'Deploy static frontend files with bundled offline dependencies'),
        ('Database Migration', 'Export existing data and import into new database'),
        ('Configuration', 'Configure all services for internal network operation'),
        ('User Migration', 'Recreate user accounts and role assignments'),
        ('SSL Setup', 'Configure HTTPS using client-provided or self-signed certificates'),
        ('Backup Configuration', 'Set up automated daily database backups'),
        ('Testing', 'Verify all functionality works correctly post-migration'),
        ('Documentation', 'Provide system administration guide'),
        ('Training', 'Remote training session for IT staff (up to 2 hours)'),
        ('Support', '30 days post-migration support'),
    ]
    table = doc.add_table(rows=len(in_scope)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Item'
    table.rows[0].cells[1].text = 'Description'
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (item, desc) in enumerate(in_scope, 1):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph('')
    doc.add_heading('2.2 Out of Scope', level=2)
    out_scope = [
        'Server hardware procurement (client responsibility)',
        'Operating system installation (client responsibility)',
        'Network/firewall configuration (client IT responsibility)',
        'Active Directory/LDAP integration (can be quoted separately)',
        'Custom feature development (separate engagement)',
        'On-site work (remote deployment; on-site available at additional cost)',
    ]
    for item in out_scope:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 3: Deliverables
    doc.add_heading('3. Deliverables', level=1)
    deliverables = [
        'Docker Compose configuration files',
        'Frontend application files (offline-ready)',
        'Database schema and seed data',
        'Environment configuration template',
        'API gateway and reverse proxy configurations',
        'Backup script',
        'System Administration Guide',
        'Deployed and tested system on client infrastructure',
        'User accounts configured with appropriate roles',
    ]
    for i, d in enumerate(deliverables, 1):
        doc.add_paragraph(f'{i}. {d}')
    
    # Section 4: Client Responsibilities
    doc.add_heading('4. Client Responsibilities', level=1)
    doc.add_paragraph('Glass Aero agrees to provide:')
    client_resp = [
        'Server meeting minimum requirements (4 cores, 8GB RAM, 100GB SSD)',
        'Operating system installed (Ubuntu 22.04 LTS or Windows Server 2019+)',
        'Remote access (SSH/RDP) to server for deployment',
        'Docker and Docker Compose installed',
        'Internal DNS entry or hosts file configuration',
        'Firewall rules allowing ports 80/443 from internal network',
        'List of users with email addresses and role assignments',
        'Designated IT contact available during deployment',
        'End-user testing and sign-off within acceptance period',
    ]
    for item in client_resp:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 5: Timeline
    doc.add_heading('5. Timeline', level=1)
    
    timeline = [
        ('Phase', 'Duration', 'Activities'),
        ('Preparation', '1-2 days', 'Export data, prepare configs, bundle dependencies'),
        ('Deployment', '1 day', 'Deploy to server, configure services, import data'),
        ('Testing', '1-2 days', 'Functional testing, user acceptance testing'),
        ('Training & Handoff', '1 day', 'IT training, documentation review, go-live'),
        ('Support', '30 days', 'Post-migration support period'),
    ]
    table = doc.add_table(rows=len(timeline), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(timeline):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Total Project Duration: ').bold = True
    p.add_run('5-7 business days (excluding support period)')
    
    # Section 6: Acceptance Criteria
    doc.add_heading('6. Acceptance Criteria', level=1)
    doc.add_paragraph('The project will be considered complete when:')
    criteria = [
        'Application accessible from internal URL',
        'Users can log in with migrated credentials',
        'All existing data migrated (repair orders, inventory, BOM)',
        'Core functions verified (create order, advance stages, BOM issuance, dashboard, reports)',
        'System functions with no internet connectivity',
        'Automated backup script runs successfully',
        'Client IT staff trained on administration',
    ]
    for i, c in enumerate(criteria, 1):
        doc.add_paragraph(f'{i}. {c}')
    
    # Section 7: Pricing
    doc.add_heading('7. Pricing', level=1)
    
    doc.add_heading('7.1 Migration Services', level=2)
    p = doc.add_paragraph()
    p.add_run('On-Premises Migration (as described in Scope): ').bold = True
    p.add_run('$2,500')
    
    doc.add_paragraph('')
    doc.add_heading('7.2 Software License Options', level=2)
    
    license_options = [
        ('Option', 'Amount'),
        ('Option A: One-Time License', '$13,000'),
        ('Option B: License + Support Retainer', '$13,000 + $500/month'),
        ('Option C: License + Hourly Support', '$13,000 + $75/hour as needed'),
    ]
    table = doc.add_table(rows=len(license_options), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(license_options):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    doc.add_heading('7.3 Total Investment (Option A)', level=2)
    
    total = [
        ('Component', 'Amount'),
        ('Software License', '$13,000'),
        ('Migration Services', '$2,500'),
        ('Total', '$15,500'),
    ]
    table = doc.add_table(rows=len(total), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(total):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0 or i == len(total)-1:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    doc.add_heading('7.4 Payment Terms', level=2)
    
    payments = [
        ('Milestone', 'Amount', 'Due'),
        ('License Fee', '$13,000', 'Upon signing'),
        ('Migration Fee', '$2,500', 'Upon go-live acceptance'),
    ]
    table = doc.add_table(rows=len(payments), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(payments):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # Section 8: Support
    doc.add_heading('8. Support', level=1)
    
    doc.add_heading('8.1 Included Support (30 Days)', level=2)
    support_included = [
        'Bug fixes for migration-related issues',
        'Configuration adjustments',
        'Assistance with user account management',
        'Email support with 24-hour response time',
        'Up to 2 hours of remote troubleshooting calls',
    ]
    for item in support_included:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('8.2 Ongoing Support Options', level=2)
    ongoing = [
        ('Monthly Retainer', '$500/month — priority support, up to 5 hours of modifications'),
        ('Hourly', '$75/hour — billed monthly for work performed'),
    ]
    for label, desc in ongoing:
        p = doc.add_paragraph()
        p.add_run(label + ': ').bold = True
        p.add_run(desc)
    
    # Section 9: Signatures
    doc.add_heading('9. Signatures', level=1)
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('JCD Enterprises').bold = True
    doc.add_paragraph('Name: _______________________________')
    doc.add_paragraph('Signature: _______________________________')
    doc.add_paragraph('Date: _______________________________')
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Glass Aero').bold = True
    doc.add_paragraph('Name: _______________________________')
    doc.add_paragraph('Title: _______________________________')
    doc.add_paragraph('Signature: _______________________________')
    doc.add_paragraph('Date: _______________________________')
    
    doc.add_paragraph('')
    p = doc.add_paragraph('This Statement of Work is subject to the terms of the Software License Agreement between the parties.')
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(9)
    
    return doc


def create_onprem_deployment_guide():
    """Create on-premises deployment guide for IT staff"""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Glass Aero Production Tracker', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('On-Premises Deployment Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_paragraph('')
    
    # Overview
    doc.add_heading('Deployment Overview', level=1)
    doc.add_paragraph(
        'This guide covers deploying the Production Tracker on internal infrastructure '
        'with no external dependencies. Once deployed, the system runs entirely within '
        'your network with no internet connectivity required.'
    )
    
    # Components
    doc.add_heading('Components to Deploy', level=2)
    components = [
        ('Component', 'Technology', 'Purpose'),
        ('Frontend', 'Static HTML/JS/CSS', 'User interface (web pages)'),
        ('Database', 'PostgreSQL 15+', 'All application data'),
        ('API Layer', 'PostgREST', 'REST API for the database'),
        ('Authentication', 'GoTrue', 'User login and sessions'),
        ('Realtime', 'Supabase Realtime', 'Live updates across browsers'),
        ('File Storage', 'Supabase Storage', 'Document uploads'),
    ]
    table = doc.add_table(rows=len(components), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(components):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # Deployment Options
    doc.add_heading('Deployment Options', level=1)
    
    options = [
        ('Option', 'Complexity', 'Best For'),
        ('A. Docker Compose', 'Low', 'Quick setup, easy updates (recommended)'),
        ('B. Manual Install', 'High', 'Maximum control, no Docker dependency'),
    ]
    table = doc.add_table(rows=len(options), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(options):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Recommendation: ').bold = True
    p.add_run('Option A (Docker) unless your organization prohibits containers.')
    
    # Server Requirements
    doc.add_heading('Server Requirements', level=1)
    
    reqs = [
        ('Resource', 'Minimum', 'Recommended'),
        ('CPU', '2 cores', '4 cores'),
        ('RAM', '4 GB', '8 GB'),
        ('Disk', '50 GB SSD', '100 GB SSD'),
        ('OS', 'Ubuntu 22.04 LTS, Windows Server 2019+, or any Docker-compatible OS', ''),
    ]
    table = doc.add_table(rows=len(reqs), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(reqs):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # Network Requirements
    doc.add_heading('Network Requirements', level=2)
    
    ports = [
        ('Port', 'Service', 'Notes'),
        ('80', 'HTTP', 'Redirect to HTTPS (optional)'),
        ('443', 'HTTPS', 'Main application access'),
        ('5432', 'PostgreSQL', 'Only if external DB access needed'),
    ]
    table = doc.add_table(rows=len(ports), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(ports):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if i == 0:
                for run in table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    # High-level steps
    doc.add_heading('Docker Deployment Steps', level=1)
    
    steps = [
        'Install Docker and Docker Compose on the server',
        'Create project directory structure',
        'Configure docker-compose.yml with all services',
        'Create environment file with passwords and secrets',
        'Configure Kong API gateway routing',
        'Configure nginx for frontend and reverse proxy',
        'Copy frontend files (HTML, JS, CSS)',
        'Update supabase-config.js with internal server URL',
        'Copy SQL schema files',
        'Start the Docker stack',
        'Create initial admin user',
        'Configure DNS or hosts file for internal access',
        'Set up SSL certificates (optional for internal network)',
        'Configure backup schedule',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    # Offline Dependencies
    doc.add_heading('Bundling for Offline Use', level=1)
    doc.add_paragraph(
        'The current frontend loads Tailwind CSS, Supabase JS, and JsBarcode from CDN. '
        'For fully offline operation, these must be downloaded and included locally:'
    )
    
    deps = [
        'Tailwind CSS (tailwind.min.css)',
        'Supabase JS client (supabase.min.js)',
        'JsBarcode (JsBarcode.all.min.js)',
    ]
    for d in deps:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_paragraph(
        'HTML files will be updated to reference local copies instead of CDN URLs.'
    )
    
    # Migration Checklist
    doc.add_heading('Migration Checklist', level=1)
    
    doc.add_heading('Pre-Migration (Developer Site)', level=2)
    pre_items = [
        'Export current database',
        'Document all user accounts and roles',
        'Package frontend files',
        'Download CDN dependencies for offline use',
        'Test deployment on staging server',
    ]
    for item in pre_items:
        doc.add_paragraph(f'☐ {item}')
    
    doc.add_heading('Client Site Deployment', level=2)
    deploy_items = [
        'Provision server (VM or physical)',
        'Install Docker and Docker Compose',
        'Create directory structure and copy config files',
        'Copy frontend files with offline dependencies',
        'Import database schema and data',
        'Start Docker stack',
        'Create admin user account',
        'Configure DNS or hosts file',
        'Test all functionality',
        'Configure SSL (if required)',
        'Set up backup schedule',
    ]
    for item in deploy_items:
        doc.add_paragraph(f'☐ {item}')
    
    doc.add_heading('Post-Migration', level=2)
    post_items = [
        'Train users on new URL',
        'Verify all reports work',
        'Confirm document uploads work',
        'Test barcode scanning',
        'Monitor logs for first week',
    ]
    for item in post_items:
        doc.add_paragraph(f'☐ {item}')
    
    # Backup
    doc.add_heading('Backup and Recovery', level=1)
    doc.add_paragraph(
        'A scheduled backup script should be configured to run daily. This will export '
        'the PostgreSQL database and should be stored on a separate drive or network share. '
        'The file storage folder (uploaded documents) should also be included in backups.'
    )
    
    # Support
    doc.add_heading('Support', level=1)
    doc.add_paragraph('30 days of post-migration support is included with deployment.')
    
    # Footer
    doc.add_paragraph('')
    p = doc.add_paragraph('For detailed configuration files and commands, see the full Deployment Guide (Markdown version).')
    p.runs[0].font.italic = True
    
    doc.add_paragraph('')
    p = doc.add_paragraph('April 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    return doc


if __name__ == '__main__':
    import os
    
    base_path = os.path.dirname(os.path.abspath(__file__))
    
    # Generate Executive Summary
    exec_doc = create_executive_summary()
    exec_path = os.path.join(base_path, 'Glass Aero Production Tracker - IT Overview.docx')
    exec_doc.save(exec_path)
    print(f'Created: {exec_path}')
    
    # Generate Technical Specification
    tech_doc = create_technical_specification()
    tech_path = os.path.join(base_path, 'Glass Aero Production Tracker - Technical Specification.docx')
    tech_doc.save(tech_path)
    print(f'Created: {tech_path}')
    
    # Generate On-Premises Deployment Guide
    deploy_doc = create_onprem_deployment_guide()
    deploy_path = os.path.join(base_path, 'Glass Aero Production Tracker - On-Premises Deployment Guide.docx')
    deploy_doc.save(deploy_path)
    print(f'Created: {deploy_path}')
    
    # Generate Migration Statement of Work
    sow_doc = create_migration_sow()
    sow_path = os.path.join(base_path, 'Glass Aero Production Tracker - Migration Statement of Work.docx')
    sow_doc.save(sow_path)
    print(f'Created: {sow_path}')
