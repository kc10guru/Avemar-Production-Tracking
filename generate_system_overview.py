"""
Generate a comprehensive Word document overview of the Glass Aero Production Tracking System.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def add_heading_style(doc):
    """Configure document styles"""
    style = doc.styles['Heading 1']
    font = style.font
    font.size = Pt(16)
    font.bold = True
    font.color.rgb = RGBColor(0x0a, 0x16, 0x28)
    
    style = doc.styles['Heading 2']
    font = style.font
    font.size = Pt(14)
    font.bold = True
    font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    
    doc.add_paragraph()

def create_overview_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Glass Aero Production Tracking System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Technical Overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Document Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Prepared by: JCD Enterprises')
    doc.add_paragraph()
    
    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'The Glass Aero Production Tracking System is a web-based application designed to track '
        'windshield repair orders through a 15-stage production workflow. The system provides '
        'real-time visibility into production status, inventory management, bill of materials '
        'tracking, and comprehensive reporting.'
    )
    doc.add_paragraph(
        'The system runs entirely on-premises on the client\'s virtual machine (VM) with no '
        'external internet dependencies. All data stays within the client\'s network.'
    )
    
    # Key capabilities
    doc.add_heading('Key Capabilities', level=2)
    bullets = [
        '15-stage production pipeline with real-time dashboard',
        'Repair order lifecycle management (create, edit, advance, revert, hold/resume)',
        'Inspection stage with customizable repair path selection',
        'Bill of Materials (BOM) with automatic parts issuance per stage',
        'Inventory management with projected needs calculations',
        'Document management with file uploads',
        'Barcode label printing for work order tracking',
        'Role-based access (Standard User vs Admin)',
        'Real-time updates across all connected browsers',
        'Business hours configuration for accurate stage timing',
        'Excel import for bulk repair order creation',
        'Weekly, Quarterly, and Annual production reports'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # =========================================================================
    # SYSTEM ARCHITECTURE
    # =========================================================================
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_heading('Architecture Overview', level=2)
    doc.add_paragraph(
        'The system uses a static frontend architecture with a Supabase-compatible backend stack. '
        'There is no custom application server - the frontend JavaScript communicates directly '
        'with the database layer through REST APIs and WebSockets.'
    )
    
    # Architecture diagram (text-based)
    doc.add_paragraph('Architecture Flow:', style='Intense Quote')
    arch_text = """
    Client Browsers (Desktop, Tablet, Mobile)
              │
              │ Port 8080
              ▼
    ┌─────────────────────────────────────┐
    │   nginx (Frontend Web Server)       │
    │   - Serves HTML/JS/CSS files        │
    │   - Proxies API calls to Kong       │
    └─────────────────┬───────────────────┘
                      │
                      ▼
    ┌─────────────────────────────────────┐
    │   Kong (API Gateway)                │
    │   - Routes: /rest, /auth,           │
    │     /realtime, /storage             │
    └──┬──────┬──────┬──────┬─────────────┘
       │      │      │      │
       ▼      ▼      ▼      ▼
    PostgREST GoTrue Realtime Storage
    (REST API) (Auth) (WebSocket) (Files)
       │      │      │      │
       └──────┴──────┴──────┘
                  │
                  ▼
    ┌─────────────────────────────────────┐
    │   PostgreSQL Database               │
    │   (Supabase Postgres image)         │
    └─────────────────────────────────────┘
    """
    para = doc.add_paragraph()
    run = para.add_run(arch_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # Technology stack table
    doc.add_heading('Technology Stack', level=2)
    add_table(doc,
        ['Layer', 'Technology', 'Purpose'],
        [
            ['Frontend UI', 'HTML + JavaScript + Tailwind CSS', 'User interface - all web pages and interactions'],
            ['API Layer', 'PostgREST', 'Automatic REST API over PostgreSQL (/rest/v1/...)'],
            ['Authentication', 'GoTrue (Supabase Auth)', 'User login, sessions, JWT tokens (/auth/v1/...)'],
            ['Real-time', 'Supabase Realtime', 'Live updates via WebSockets (/realtime/v1/...)'],
            ['File Storage', 'Supabase Storage API', 'Document uploads (/storage/v1/...)'],
            ['API Gateway', 'Kong', 'Routes requests to correct backend services'],
            ['Web Server', 'nginx', 'Serves static files, proxies API to Kong'],
            ['Database', 'PostgreSQL 15', 'All application data, auth schema, storage metadata'],
        ]
    )
    
    # =========================================================================
    # VM DEPLOYMENT DETAILS
    # =========================================================================
    doc.add_heading('3. VM Deployment Details', level=1)
    
    doc.add_heading('Docker Containers', level=2)
    doc.add_paragraph(
        'The entire system runs as Docker containers orchestrated by Docker Compose. '
        'Seven containers work together to provide the full functionality:'
    )
    
    add_table(doc,
        ['Container Name', 'Image', 'Internal Port', 'Host Port', 'Purpose'],
        [
            ['glass-aero-db', 'supabase/postgres:15.8.1.085', '5432', '5435', 'PostgreSQL database'],
            ['glass-aero-rest', 'postgrest/postgrest:v12.0.2', '3000', '-', 'REST API'],
            ['glass-aero-auth', 'supabase/gotrue:v2.143.0', '9999', '-', 'Authentication'],
            ['glass-aero-realtime', 'supabase/realtime:v2.25.35', '4000', '-', 'WebSocket updates'],
            ['glass-aero-storage', 'supabase/storage-api:v0.43.11', '5000', '-', 'File uploads'],
            ['glass-aero-kong', 'kong:2.8.1', '8000', '8300', 'API Gateway'],
            ['glass-aero-frontend', 'nginx:alpine', '80', '8080', 'Web server'],
        ]
    )
    
    doc.add_heading('Host Ports (Firewall)', level=2)
    doc.add_paragraph('Only these ports need to be accessible:')
    add_table(doc,
        ['Port', 'Service', 'Who Uses It'],
        [
            ['8080', 'Frontend (nginx)', 'All users - this is the main URL (http://SERVER:8080)'],
            ['8300', 'Kong API Gateway', 'Optional - for direct API access or debugging'],
            ['5435', 'PostgreSQL', 'Admins only - for database backups and direct queries'],
        ]
    )
    
    doc.add_heading('File Locations on VM', level=2)
    doc.add_paragraph('Default installation path: C:\\glass-aero-tracker (Windows) or /opt/glass-aero-tracker (Linux)')
    add_table(doc,
        ['Path', 'Contents'],
        [
            ['docker-compose.yml', 'Docker services configuration'],
            ['.env', 'Environment variables (passwords, JWT secrets, site URL)'],
            ['frontend/', 'HTML, JavaScript, CSS files'],
            ['frontend/js/', 'Application JavaScript files'],
            ['frontend/vendor/', 'Third-party libraries (Supabase JS, JsBarcode, etc.)'],
            ['sql/', 'Database initialization and migration scripts'],
            ['kong/kong.yml', 'API Gateway routing configuration'],
            ['nginx/nginx.conf', 'Web server and proxy configuration'],
            ['data/postgres/', 'PostgreSQL database files (IMPORTANT: This is your data)'],
            ['data/storage/', 'Uploaded documents and files'],
            ['backups/', 'Database backup files (if backup script is configured)'],
        ]
    )
    
    # =========================================================================
    # DATABASE STRUCTURE
    # =========================================================================
    doc.add_heading('4. Database Structure', level=1)
    
    doc.add_heading('Database Location', level=2)
    doc.add_paragraph(
        'The PostgreSQL database runs inside the glass-aero-db container. The actual data files '
        'are stored on the VM\'s file system at:'
    )
    doc.add_paragraph('Windows: C:\\glass-aero-tracker\\data\\postgres\\', style='List Bullet')
    doc.add_paragraph('Linux: /opt/glass-aero-tracker/data/postgres/', style='List Bullet')
    doc.add_paragraph(
        'This folder is mounted into the Docker container, so the data persists even if the '
        'container is stopped or recreated. THIS IS THE CRITICAL DATA TO BACK UP.'
    )
    
    doc.add_heading('Main Database Tables', level=2)
    add_table(doc,
        ['Table Name', 'Purpose', 'Key Fields'],
        [
            ['repair_orders', 'All repair orders (work orders)', 'id, ro_number, customer_name, part_number, serial_number, current_stage, status, is_on_hold, skipped_stages'],
            ['production_stages', 'The 15 production workflow stages', 'id, stage_number, name, time_limit_hours'],
            ['production_parts', 'Windshield part numbers (assemblies)', 'id, part_number, description'],
            ['stage_history', 'Audit trail of stage transitions', 'repair_order_id, stage_number, entered_at, exited_at, completed_by'],
            ['subcomponents', 'Inventory items (parts/materials)', 'id, part_number, description, quantity_on_hand, reorder_point'],
            ['bom_items', 'Bill of Materials - parts needed per stage', 'production_part_id, subcomponent_id, stage_number, quantity'],
            ['parts_issuance', 'Record of parts issued to repair orders', 'repair_order_id, subcomponent_id, stage_number, quantity'],
            ['repair_order_documents', 'Metadata for uploaded files', 'repair_order_id, file_name, storage_path'],
            ['hold_history', 'Record of hold/resume events', 'repair_order_id, hold_reason, started_at, ended_at, duration_hours'],
            ['app_settings', 'Application configuration', 'key, value (business hours, timezone, etc.)'],
        ]
    )
    
    doc.add_heading('Production Stages (15 Total)', level=2)
    stages = [
        (1, 'Check-In', 'Receiving and initial logging'),
        (2, 'Pre-Inspection', 'Preliminary assessment'),
        (3, 'Inspection', 'Detailed inspection - inspector selects required stages'),
        (4, 'Remove Hardware', 'Disassembly'),
        (5, 'Remove Glass', 'Glass removal'),
        (6, 'Prep Outer Frame', 'Frame preparation'),
        (7, 'Prep Inner Frame', 'Inner frame work'),
        (8, 'Heater Installation', 'Can be skipped based on inspection'),
        (9, 'Outer Glass Installation', 'Can be skipped based on inspection'),
        (10, 'Inner Glass Installation', 'Glass installation'),
        (11, 'Install Hardware', 'Reassembly'),
        (12, 'Seal and Cure', 'Sealing process'),
        (13, 'Final Inspection', 'Quality check'),
        (14, 'Shipping', 'Packaging'),
        (15, 'Delivered', 'Completion'),
    ]
    add_table(doc, ['Stage #', 'Name', 'Description'], stages)
    
    # =========================================================================
    # WEB APPLICATION FUNCTIONALITY
    # =========================================================================
    doc.add_heading('5. Web Application Functionality', level=1)
    
    doc.add_heading('Pages and Features', level=2)
    
    pages = [
        ('login.html', 'User authentication - email/password login'),
        ('index.html', 'Dashboard - production overview, pipeline visualization, stats, alerts'),
        ('repair-orders.html', 'List all repair orders with search, filter, sort'),
        ('new-repair-order.html', 'Create new repair order form'),
        ('repair-order-detail.html', 'Full order details, stage history, documents, parts, advance/hold'),
        ('scan.html', 'Barcode scanning page for quick stage advancement'),
        ('inventory.html', 'Manage subcomponent stock levels'),
        ('bom.html', 'Configure Bill of Materials per part number (Admin)'),
        ('reports.html', 'Weekly, Quarterly, Annual production reports'),
        ('import.html', 'Bulk import repair orders from Excel'),
        ('settings.html', 'System configuration - parts, stages, business hours (Admin)'),
    ]
    add_table(doc, ['Page', 'Purpose'], pages)
    
    doc.add_heading('Dashboard Features', level=2)
    bullets = [
        'Active Orders count - total repair orders in progress',
        'Late Units - orders that exceeded stage time limits (based on business hours)',
        'On Hold count - orders currently paused',
        'Completed This Month - finished orders',
        'Low Stock Alerts - inventory items below reorder point',
        'Production Pipeline - visual grid showing unit counts at each stage',
        'Real-time updates - changes appear instantly across all browsers',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Repair Order Workflow', level=2)
    doc.add_paragraph('Each repair order flows through these actions:')
    bullets = [
        'Create: New order with customer info, part number, serial number',
        'Edit: Modify any field including part number (with BOM correction)',
        'Advance Stage: Move to next stage, auto-issue BOM parts',
        'Go Back: Revert to previous stage, restore issued parts to inventory',
        'Hold: Pause work with required reason',
        'Resume: Continue work after hold',
        'Delete: Remove order completely (with confirmation)',
        'Print Barcode: Generate scannable label',
        'Upload Documents: Attach files to the order',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Inspection Stage (Stage 3)', level=2)
    doc.add_paragraph(
        'The inspection stage is special. When advancing from Stage 3, the inspector sees a '
        'checklist of all production stages (4-13) and can uncheck any stages that the windshield '
        'doesn\'t need. For example, if a windshield only needs polishing and not heater/glass work, '
        'the inspector unchecks stages 8 and 9.'
    )
    doc.add_paragraph('Effects of skipping stages:')
    bullets = [
        'Skipped stages appear dashed/faded in the progress bar',
        'Advancing automatically jumps over skipped stages',
        'BOM parts for skipped stages are never issued',
        'Projected inventory excludes parts from skipped stages',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Scan Page', level=2)
    doc.add_paragraph(
        'The scan page is designed for shop floor use with barcode scanners or tablets. '
        'It provides a simplified interface for quick stage advancement without needing the full detail page.'
    )
    bullets = [
        'Large input field for scanner-as-keyboard input',
        'Shows current stage and unit info',
        'Big "Advance" button for touch screens',
        'Works with USB or Bluetooth barcode scanners',
        'For inspection stage (Stage 3), redirects to full detail page',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Business Hours', level=2)
    doc.add_paragraph(
        'Stage time limits are calculated using business hours only. The system does not count '
        'nights, weekends, or holidays toward stage time. Business hours are configurable in Settings:'
    )
    bullets = [
        'Shop Opens / Shop Closes times (default: 9 AM - 6 PM)',
        'Timezone selection (Eastern, Central, Mountain, Pacific, etc.)',
        'Work Days checkboxes (default: Monday - Friday)',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # =========================================================================
    # DATA FLOW AND INTERACTIONS
    # =========================================================================
    doc.add_heading('6. Data Flow and Interactions', level=1)
    
    doc.add_heading('How Components Interact', level=2)
    
    flows = [
        ('User opens browser', 'nginx serves HTML/JS/CSS files from frontend/ folder'),
        ('User logs in', 'JS calls /auth/v1/ → Kong routes to GoTrue → validates against auth schema'),
        ('User views dashboard', 'JS calls /rest/v1/repair_orders → Kong routes to PostgREST → queries PostgreSQL'),
        ('User advances stage', 'JS calls /rest/v1/repair_orders (PATCH) → updates database → triggers realtime'),
        ('Another user sees update', 'Realtime service pushes change via WebSocket → browser updates automatically'),
        ('User uploads document', 'JS calls /storage/v1/ → Kong routes to Storage API → saves to data/storage/'),
        ('User prints barcode', 'JsBarcode library generates Code128 barcode in browser → sends to printer'),
    ]
    add_table(doc, ['User Action', 'System Response'], flows)
    
    doc.add_heading('Bill of Materials (BOM) Flow', level=2)
    doc.add_paragraph('When a repair order advances to a new stage:')
    bullets = [
        '1. System looks up BOM items for that part number and stage',
        '2. If BOM items exist, they are displayed in the advance modal',
        '3. User confirms the advance',
        '4. System decrements quantity_on_hand for each subcomponent',
        '5. System creates parts_issuance records',
        '6. Dashboard projected inventory recalculates automatically',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_heading('Part Number Change Flow', level=2)
    doc.add_paragraph('When editing a repair order and changing the part number:')
    bullets = [
        '1. System reverses all previously issued BOM parts (restores inventory)',
        '2. System deletes old parts_issuance records',
        '3. System looks up new BOM for the new part number',
        '4. System re-issues correct BOM parts for all completed stages',
        '5. Projected inventory automatically corrects',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    # =========================================================================
    # BACKUP AND RECOVERY
    # =========================================================================
    doc.add_heading('7. Backup and Recovery', level=1)
    
    doc.add_heading('Critical Data to Back Up', level=2)
    add_table(doc,
        ['Location', 'Contents', 'Priority'],
        [
            ['data/postgres/', 'PostgreSQL database files - ALL application data', 'CRITICAL'],
            ['data/storage/', 'Uploaded documents and files', 'HIGH'],
            ['.env', 'Environment configuration (passwords, secrets)', 'HIGH'],
            ['frontend/', 'Application files (can be restored from source)', 'LOW'],
        ]
    )
    
    doc.add_heading('Database Backup Command', level=2)
    doc.add_paragraph('To create a database backup:')
    para = doc.add_paragraph()
    run = para.add_run('docker exec glass-aero-db pg_dump -U postgres postgres > backup.sql')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    
    doc.add_heading('Database Restore Command', level=2)
    doc.add_paragraph('To restore from a backup:')
    para = doc.add_paragraph()
    run = para.add_run('cat backup.sql | docker exec -i glass-aero-db psql -U postgres postgres')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    
    # =========================================================================
    # TROUBLESHOOTING
    # =========================================================================
    doc.add_heading('8. Quick Troubleshooting', level=1)
    
    add_table(doc,
        ['Issue', 'Likely Cause', 'Solution'],
        [
            ['Page won\'t load', 'Container stopped', 'Run: docker compose ps (check all running)'],
            ['Login fails', 'Auth container issue', 'Run: docker compose logs auth'],
            ['Changes not appearing', 'Browser cache', 'Hard refresh: Ctrl+Shift+R'],
            ['Real-time not working', 'WebSocket blocked', 'Check nginx config and firewall'],
            ['Storage uploads fail', 'Storage container', 'Run: docker compose logs storage'],
            ['Database connection error', 'DB container issue', 'Run: docker compose logs db'],
        ]
    )
    
    doc.add_heading('Useful Commands', level=2)
    commands = [
        ('Check all containers', 'docker compose ps'),
        ('View all logs', 'docker compose logs -f'),
        ('Restart all services', 'docker compose restart'),
        ('Stop everything', 'docker compose down'),
        ('Start everything', 'docker compose up -d'),
        ('Connect to database', 'docker exec -it glass-aero-db psql -U postgres postgres'),
    ]
    add_table(doc, ['Action', 'Command'], commands)
    
    # =========================================================================
    # USER ROLES
    # =========================================================================
    doc.add_heading('9. User Roles and Permissions', level=1)
    
    add_table(doc,
        ['Role', 'Capabilities'],
        [
            ['Standard User', 'View dashboard, create/edit orders, advance stages, manage inventory, view reports'],
            ['Admin', 'All standard user capabilities PLUS: BOM configuration, Settings, Import, Delete orders, manage part numbers, edit stage time limits, configure business hours'],
        ]
    )
    
    doc.add_paragraph(
        'User roles are set in the database via app_metadata. To make a user an admin:'
    )
    para = doc.add_paragraph()
    cmd = 'docker exec -i glass-aero-db psql -U postgres -d postgres -c "UPDATE auth.users SET raw_app_meta_data = raw_app_meta_data || \'{\\"role\\": \\"admin\\"}\' WHERE email = \'user@example.com\';"'
    run = para.add_run(cmd)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    # =========================================================================
    # SUPPORT INFORMATION
    # =========================================================================
    doc.add_heading('10. Support Information', level=1)
    doc.add_paragraph('Developer: JCD Enterprises')
    doc.add_paragraph('System Version: 2.0')
    doc.add_paragraph(f'Document Generated: {datetime.now().strftime("%B %d, %Y")}')
    
    # Save document
    output_path = r'c:\Users\Jeff\OneDrive\Documents\JCD Enterprises\Avemar-Production-Tracking\Glass Aero Production Tracker - System Overview.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_overview_doc()
